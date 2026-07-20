#!/usr/bin/env python3
"""
Declaration Form Generator
===========================
A single-file desktop app (Tkinter GUI) that generates the VIT
"Student Undertaking for Ethical Academic Practice" declaration form
as a Word (.docx) and/or PDF file, for any subject/vertical and any
assignment/experiment number.

This version is built to match the reference form pixel-for-pixel:
    - The real "Quattrocento Sans" font is EMBEDDED directly inside
      the generated .docx (copied from a "fonts" folder next to this
      script), so it renders correctly on any machine even if the
      font isn't installed -- this is what was missing before and is
      why the old output looked like it fell back to a random default
      font.
    - The PDF uses the same embedded .ttf files via reportlab, so the
      .docx and the .pdf are visually identical.
    - Two-column fields (Branch/Vertical, Semester/Division, Subject
      Name/Code, Name/Roll No, Signature/Date) use real tab stops in
      the DOCX and a real fixed-width table in the PDF -- not
      hand-counted spaces -- so columns line up regardless of how
      long the text in them is.
    - Paragraph spacing (8pt after / 1.1583 line spacing), 10pt body
      size, bold-for-values / regular-for-labels, left-aligned bold
      titles, and a justified ethics-warning paragraph all match the
      reference form exactly.
    - The signature image pipeline was rewritten. The old version
      over-thresholded the image and frequently rendered as a solid
      black block. It now auto-contrasts, crops to the ink, and
      anti-aliases the result onto a clean white background at the
      same size used in the reference form.

Requirements (install once):
    pip install python-docx reportlab pillow

Run:
    python declaration_form_generator.py

Output location:
    Files are saved automatically into "forms/docs" (.docx) and
    "forms/pdfs" (.pdf), created next to this script if they don't
    already exist. File name pattern:
    <SubjectCode>-<AssignmentNo>-Declaration Form.docx / .pdf
    e.g.  forms/docs/PCCE10T-Experiment1-Declaration Form.docx

Fonts (REQUIRED for a pixel-accurate match):
    Put these four files in a "fonts" folder next to this script:
        QuattrocentoSans-regular.ttf
        QuattrocentoSans-bold.ttf
        QuattrocentoSans-italic.ttf
        QuattrocentoSans-boldItalic.ttf
    They are bundled together with this script. If they are missing,
    the DOCX still *requests* "Quattrocento Sans" (Word will
    substitute a similar font on machines that don't have it), and
    the PDF quietly falls back to Helvetica.

Optional: drop a "college_logo.png" file next to this script and it
will be embedded at the top of both the DOCX and PDF outputs. You can
also attach a signature image (PNG/JPG) from the app's "Signature
Image" field -- it gets cleaned up and embedded on the Signature line
of both files.
"""

import os
import io
import sys
import shutil
import zipfile
import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# --------------------------------------------------------------------------
# 0. OUTPUT FOLDERS  (created automatically next to this script)
# --------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_DIR = os.path.join(SCRIPT_DIR, "forms", "docs")
PDF_DIR = os.path.join(SCRIPT_DIR, "forms", "pdfs")
FONT_DIR = os.path.join(SCRIPT_DIR, "fonts")
os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

# --------------------------------------------------------------------------
# 0.5 DESIGN CONSTANTS  (lifted from the reference form)
# --------------------------------------------------------------------------
FONT_NAME = "Quattrocento Sans"
BODY_SIZE_PT = 10                  # w:sz val="20" (half-points) -> 10pt
PARA_SPACE_AFTER_PT = 8            # w:spacing w:after="160" (twips) -> 8pt
LINE_SPACING = 278.0 / 240.0       # w:line="278" auto -> 278/240 = 1.15833
SECOND_COL_TAB_IN = 3.5            # tab stop for the 2nd field column
SIGNATURE_MAX_W_IN = 719394 / 914400.0   # from reference form's embedded size
SIGNATURE_MAX_H_IN = 295465 / 914400.0   # (0.7867in x 0.3231in)
CHECK_MARK = "\u2713"               # checked
BOX_EMPTY = "\u2610"                # unchecked

FONT_FILES = {
    "regular": os.path.join(FONT_DIR, "QuattrocentoSans-regular.ttf"),
    "bold": os.path.join(FONT_DIR, "QuattrocentoSans-bold.ttf"),
    "italic": os.path.join(FONT_DIR, "QuattrocentoSans-italic.ttf"),
    "boldItalic": os.path.join(FONT_DIR, "QuattrocentoSans-boldItalic.ttf"),
}

# --------------------------------------------------------------------------
# 1. VERTICAL / SUBJECT DATA
#    Edit this dictionary to add / change verticals, subjects or codes.
#    Format:  "Vertical Key": [("Subject Name", "Subject Code"), ...]
# --------------------------------------------------------------------------
VERTICALS = {
    "HSSM_VEC": [
        ("E-Waste and Environmental Management", "VEC02"),
    ],
    "MDC_MDM": [
        ("Intelligent Mobile Robotics", "MDMRB03"),
        ("Innovation Management and Scaling Startups", "MDMIE03"),
        ("Machine Learning Applications in Bioinfomatics", "MDMBI03"),
        ("Strategic Marketing and Business Planning", "MDMBI03"),
    ],
    "MDC_OEC": [
        ("Professional Competency Development", "OEC15"),
    ],
    "PC_PCC": [
        ("Artificial Intelligence", "PCCE10T"),
        ("Computer Networks", "PCCE11T"),
        ("Software Engineering", "PCCE12T"),
        ("Theory Of Computation", "PCCE09"),
    ],
    "PC_PEC": [
        ("Data Warehousing and Mining", "PECE02T"),
    ],
}

DECLARATION_INTRO = (
    "I hereby declare that for the assignment / academic activity "
    "submitted by me, I confirm ONE of the following statements by "
    "ticking (\u2610) the appropriate box:"
)

DECLARATION_OPTIONS = [
    "I have prepared the assignment / academic activity entirely by myself.",
    "I have completed the assignment / academic activity with academic "
    "guidance or discussion from seniors / friends / peers, while ensuring "
    "the work is my own.",
    "I have used AI-based tools only as an aid and have appropriately "
    "acknowledged or cited their use as per the ethical guidelines of the "
    "Institute.",
    "I have used AI-based tools and directly copy-pasted content for the "
    "assignment / academic activity.",
    "I have copied the assignment / academic activity from peers / friends.",
]

ETHICS_WARNING = (
    "I understand that selecting options (4) or (5) indicates unethical "
    "academic practice and may attract academic penalties, including "
    "rejection of the assignment or disciplinary action, as per the "
    "rules of Vidyalankar Institute of Technology, Mumbai. I submit "
    "this declaration truthfully and accept full responsibility for "
    "the same."
)

LOGO_FILE = os.path.join(SCRIPT_DIR, "college_logo.png")


# --------------------------------------------------------------------------
# 1.5 SIGNATURE PROCESSOR
#     Auto-contrasts, crops to the ink, anti-aliases, and places the
#     result on a clean white background at a fixed max size (matches
#     the reference form's embedded signature: ~0.79in x 0.32in).
# --------------------------------------------------------------------------
def process_signature_image(img_path_or_bytes):
    """
    Returns (processed_bytes_io, target_width_inches, target_height_inches)
    or (None, 0, 0) if no image / processing failed.
    """
    if not img_path_or_bytes:
        return None, 0, 0
    try:
        from PIL import Image, ImageOps

        if isinstance(img_path_or_bytes, str):
            if not os.path.exists(img_path_or_bytes):
                return None, 0, 0
            img = Image.open(img_path_or_bytes)
        else:
            img_path_or_bytes.seek(0)
            img = Image.open(img_path_or_bytes)

        # 1. Flatten any transparency onto a plain white background --
        #    matches the reference form's signature, which is a plain
        #    white-background PNG, not a transparent cut-out.
        if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
            img = img.convert("RGBA")
            bg = Image.new("RGBA", img.size, (255, 255, 255, 255))
            img = Image.alpha_composite(bg, img).convert("RGB")
        else:
            img = img.convert("RGB")

        # 2. Grayscale + autocontrast to normalize scan/photo lighting.
        gray = img.convert("L")
        gray = ImageOps.autocontrast(gray, cutoff=1)

        # 3. Soft threshold: only true background noise is dropped to
        #    white; ink stays as its natural gray levels so edges stay
        #    smooth instead of collapsing into one solid black shape.
        gray = gray.point(lambda p: 255 if p > 235 else p)

        # 4. Crop tightly to the ink (non-white pixels), with a little
        #    padding so strokes aren't clipped.
        inverted = ImageOps.invert(gray)
        bbox = inverted.getbbox()
        if bbox:
            pad = 6
            left = max(0, bbox[0] - pad)
            top = max(0, bbox[1] - pad)
            right = min(gray.width, bbox[2] + pad)
            bottom = min(gray.height, bbox[3] + pad)
            gray = gray.crop((left, top, right, bottom))

        w, h = gray.size
        if w == 0 or h == 0:
            return None, 0, 0
        aspect = w / float(h)

        # 5. Fit inside the reference form's signature box.
        target_h = SIGNATURE_MAX_H_IN
        target_w = target_h * aspect
        if target_w > SIGNATURE_MAX_W_IN:
            target_w = SIGNATURE_MAX_W_IN
            target_h = target_w / aspect

        # 6. Render at a crisp resolution (300dpi) with Lanczos
        #    resampling for smooth, anti-aliased strokes.
        px_w = max(1, int(round(target_w * 300)))
        px_h = max(1, int(round(target_h * 300)))
        final = gray.resize((px_w, px_h), Image.Resampling.LANCZOS).convert("RGB")

        out = io.BytesIO()
        final.save(out, format="PNG")
        out.name = "signature.png"
        out.seek(0)
        return out, target_w, target_h
    except Exception:
        return None, 0, 0


# --------------------------------------------------------------------------
# 2. DOCX GENERATION
# --------------------------------------------------------------------------
def generate_docx(data, out_path):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise RuntimeError(
            "The 'python-docx' package is not installed.\n\n"
            "Open a terminal and run:\n"
            "    pip install python-docx\n\n"
            "(If 'pip' isn't recognized, try 'python -m pip install python-docx' "
            "or 'pip3 install python-docx'.)"
        )

    doc = Document()

    # --- Page setup: A4, 1 inch margins ----------------------------------
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Base style: Quattrocento Sans 10pt, 8pt-after / 1.1583 line -----
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_SIZE_PT)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT_NAME)
    normal.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    normal.paragraph_format.line_spacing = LINE_SPACING

    def set_font(run, bold=False):
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_SIZE_PT)
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.append(rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rFonts.set(qn(attr), FONT_NAME)

    def add_para(justify=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
        p.paragraph_format.line_spacing = LINE_SPACING
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def add_run(p, text, bold=False):
        r = p.add_run(text)
        set_font(r, bold=bold)
        return r

    def add_tab(p):
        r = p.add_run("\t")
        set_font(r)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(SECOND_COL_TAB_IN), WD_TAB_ALIGNMENT.LEFT)

    def add_break(p):
        r = p.add_run()
        set_font(r)
        r.add_break()

    # Optional logo
    if os.path.exists(LOGO_FILE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(LOGO_FILE, width=Inches(1.2))

    # --- Title block: left-aligned, bold, same 10pt body size -----------
    p = add_para()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Student Undertaking for Ethical Academic Practice", bold=True)

    p = add_para()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Vidyalankar Institute of Technology, Mumbai", bold=True)

    add_para()  # blank spacer line, as in the reference form

    from docx.oxml import OxmlElement

    def remove_table_borders(table):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            tblBorders = tblPr[0].xpath('w:tblBorders')
            if tblBorders:
                tblPr[0].remove(tblBorders[0])
            borders = OxmlElement('w:tblBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'none')
                borders.append(node)
            tblPr[0].append(borders)

    def add_docx_2col_table(rows_tuples, col1_in=3.5, col2_in=2.77):
        table = doc.add_table(rows=len(rows_tuples), cols=2)
        remove_table_borders(table)
        for i, row in enumerate(rows_tuples):
            c0 = table.cell(i, 0)
            c0.width = Inches(col1_in)
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p0.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
            p0.paragraph_format.line_spacing = LINE_SPACING
            left_data = row[0]
            if left_data:
                lbl, val = left_data
                add_run(p0, lbl, bold=False)
                if val:
                    add_run(p0, val, bold=True)

            c1 = table.cell(i, 1)
            c1.width = Inches(col2_in)
            p1 = c1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p1.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
            p1.paragraph_format.line_spacing = LINE_SPACING
            right_data = row[1] if len(row) > 1 else None
            if right_data:
                lbl, val = right_data
                add_run(p1, lbl, bold=False)
                if val:
                    add_run(p1, val, bold=True)
        return table

    # --- Metadata Borderless Table ---
    metadata_rows = [
        (("Assignment No: ", data["assignment_no"]), None),
        (("Branch: ", data["branch"]), ("Vertical: ", data["vertical"])),
        (("Semester: ", data["semester"]), ("Division: ", data["division"])),
        (("Subject Name: ", data["subject_name"]), ("Subject Code: ", data["subject_code"])),
    ]
    add_docx_2col_table(metadata_rows)

    # --- Declaration intro line -----------------------------------------
    p = add_para()
    add_run(p, DECLARATION_INTRO)

    # --- Declaration options: one paragraph, line breaks between items --
    p = add_para()
    for idx, text in enumerate(DECLARATION_OPTIONS, start=1):
        mark = CHECK_MARK if idx in data["selected_options"] else BOX_EMPTY
        if idx != 1:
            add_break(p)
        add_run(p, f"{mark} ({idx}) {text}")

    # --- Ethics warning: justified ---------------------------------------
    p = add_para(justify=True)
    add_run(p, ETHICS_WARNING)

    # --- Student Info & Signature Borderless Table -----------------------
    sig_table = doc.add_table(rows=2, cols=2)
    remove_table_borders(sig_table)

    # Row 0: Student Name / Roll No
    c00 = sig_table.cell(0, 0)
    c00.width = Inches(3.5)
    p00 = c00.paragraphs[0]
    p00.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p00.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    p00.paragraph_format.line_spacing = LINE_SPACING
    add_run(p00, "Student Name: ", bold=False)
    add_run(p00, data["student_name"], bold=True)

    c01 = sig_table.cell(0, 1)
    c01.width = Inches(2.77)
    p01 = c01.paragraphs[0]
    p01.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p01.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    p01.paragraph_format.line_spacing = LINE_SPACING
    add_run(p01, "Roll No.: ", bold=False)
    add_run(p01, data["roll_no"], bold=True)

    # Row 1: Signature / Date
    c10 = sig_table.cell(1, 0)
    c10.width = Inches(3.5)
    p10 = c10.paragraphs[0]
    p10.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p10.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    p10.paragraph_format.line_spacing = LINE_SPACING
    add_run(p10, "Signature: ", bold=False)
    signature_path = data.get("signature_path")
    if signature_path:
        processed_sig, tw, th = process_signature_image(signature_path)
        if processed_sig:
            sig_run = p10.add_run()
            set_font(sig_run)
            sig_run.add_picture(processed_sig, width=Inches(tw), height=Inches(th))

    c11 = sig_table.cell(1, 1)
    c11.width = Inches(2.77)
    p11 = c11.paragraphs[0]
    p11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p11.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    p11.paragraph_format.line_spacing = LINE_SPACING
    add_run(p11, "Date: ", bold=False)
    add_run(p11, data["date"], bold=True)

    doc.save(out_path)

    # Embed the real Quattrocento Sans font family into the .docx so it
    # renders correctly even on machines that don't have the font
    # installed. This is what fixes "wrong font" complaints.
    embed_fonts_in_docx(out_path)


def embed_fonts_in_docx(docx_path):
    """
    Post-processes a saved .docx to embed the Quattrocento Sans TTF
    family (from FONT_FILES) directly inside the file, the same way
    the reference form does it. Silently does nothing if the font
    files aren't present next to the script.
    """
    if not all(os.path.exists(p) for p in FONT_FILES.values()):
        return

    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        parts = {name: zin.read(name) for name in names}

    # 1. Add the four font parts.
    font_targets = {
        "regular": "word/fonts/QuattrocentoSans-regular.ttf",
        "bold": "word/fonts/QuattrocentoSans-bold.ttf",
        "italic": "word/fonts/QuattrocentoSans-italic.ttf",
        "boldItalic": "word/fonts/QuattrocentoSans-boldItalic.ttf",
    }
    for key, target in font_targets.items():
        with open(FONT_FILES[key], "rb") as f:
            parts[target] = f.read()

    # 2. fontTable.xml.rels -- relationships from fontTable.xml to the
    #    four font parts.
    rel_ids = {
        "regular": "rFontEmbed1",
        "bold": "rFontEmbed2",
        "italic": "rFontEmbed3",
        "boldItalic": "rFontEmbed4",
    }
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
    rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    )
    for key, target in font_targets.items():
        fname = target.split("/")[-1]
        rels_xml += f'<Relationship Id="{rel_ids[key]}" Type="{rel_type}" Target="fonts/{fname}"/>'
    rels_xml += "</Relationships>"
    parts["word/_rels/fontTable.xml.rels"] = rels_xml.encode("utf-8")

    # 3. fontTable.xml -- append (or create) the Quattrocento Sans entry.
    ns = (
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" mc:Ignorable="w14"'
    )
    font_entry = (
        f'<w:font w:name="{FONT_NAME}">'
        f'<w:embedRegular r:id="{rel_ids["regular"]}" w:fontKey="{{00000000-0000-0000-0000-000000000000}}" w:subsetted="0"/>'
        f'<w:embedBold r:id="{rel_ids["bold"]}" w:fontKey="{{00000000-0000-0000-0000-000000000000}}" w:subsetted="0"/>'
        f'<w:embedItalic r:id="{rel_ids["italic"]}" w:fontKey="{{00000000-0000-0000-0000-000000000000}}" w:subsetted="0"/>'
        f'<w:embedBoldItalic r:id="{rel_ids["boldItalic"]}" w:fontKey="{{00000000-0000-0000-0000-000000000000}}" w:subsetted="0"/>'
        f'</w:font>'
    )
    if "word/fontTable.xml" in parts:
        existing = parts["word/fontTable.xml"].decode("utf-8")
        insert_at = existing.rfind("</w:fonts>")
        new_xml = existing[:insert_at] + font_entry + existing[insert_at:]
    else:
        new_xml = f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n<w:fonts {ns}>{font_entry}</w:fonts>'
    parts["word/fontTable.xml"] = new_xml.encode("utf-8")

    # 4. [Content_Types].xml -- make sure .ttf has a default content type.
    ct_key = "[Content_Types].xml"
    ct_xml = parts[ct_key].decode("utf-8")
    if 'Extension="ttf"' not in ct_xml:
        insert_at = ct_xml.find("<Default")
        ct_xml = (
            ct_xml[:insert_at]
            + '<Default Extension="ttf" ContentType="application/x-font-ttf"/>'
            + ct_xml[insert_at:]
        )
        parts[ct_key] = ct_xml.encode("utf-8")

    # 5. settings.xml -- tell Word the doc embeds its own fonts.
    settings_key = "word/settings.xml"
    settings_xml = parts[settings_key].decode("utf-8")
    if "embedTrueTypeFonts" not in settings_xml:
        insert_at = settings_xml.find("<w:defaultTabStop")
        if insert_at == -1:
            insert_at = settings_xml.find("<w:compat")
        settings_xml = (
            settings_xml[:insert_at]
            + '<w:embedTrueTypeFonts w:val="1"/><w:saveSubsetFonts w:val="0"/>'
            + settings_xml[insert_at:]
        )
        parts[settings_key] = settings_xml.encode("utf-8")

    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, content in parts.items():
            zout.writestr(name, content)

    shutil.move(tmp_path, docx_path)


# --------------------------------------------------------------------------
# 3. PDF GENERATION  (reportlab -- no MS Word/LibreOffice required)
# --------------------------------------------------------------------------
def _register_pdf_fonts():
    """Register the real Quattrocento Sans family if the .ttf files are
    present in FONT_DIR; otherwise fall back to Helvetica silently."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping

    all_present = all(os.path.exists(p) for p in FONT_FILES.values())
    if not all_present:
        return "Helvetica", "Helvetica-Bold"

    pdfmetrics.registerFont(TTFont("QuattrocentoSans", FONT_FILES["regular"]))
    pdfmetrics.registerFont(TTFont("QuattrocentoSans-Bold", FONT_FILES["bold"]))
    pdfmetrics.registerFont(TTFont("QuattrocentoSans-Italic", FONT_FILES["italic"]))
    pdfmetrics.registerFont(TTFont("QuattrocentoSans-BoldItalic", FONT_FILES["boldItalic"]))
    addMapping("QuattrocentoSans", 0, 0, "QuattrocentoSans")
    addMapping("QuattrocentoSans", 1, 0, "QuattrocentoSans-Bold")
    addMapping("QuattrocentoSans", 0, 1, "QuattrocentoSans-Italic")
    addMapping("QuattrocentoSans", 1, 1, "QuattrocentoSans-BoldItalic")
    return "QuattrocentoSans", "QuattrocentoSans-Bold"


def generate_pdf(data, out_path):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
    except ImportError:
        raise RuntimeError(
            "The 'reportlab' package is not installed.\n\n"
            "Open a terminal and run:\n"
            "    pip install reportlab\n\n"
            "(If 'pip' isn't recognized, try 'python -m pip install reportlab' "
            "or 'pip3 install reportlab'.)"
        )

    base_font, bold_font = _register_pdf_fonts()
    content_width = A4[0] / inch - 2  # 1in margins each side, in inches
    col_width = SECOND_COL_TAB_IN
    leading = BODY_SIZE_PT * LINE_SPACING

    title_style = ParagraphStyle(
        "TitleL", alignment=TA_LEFT, fontName=bold_font, fontSize=BODY_SIZE_PT,
        leading=leading, spaceAfter=PARA_SPACE_AFTER_PT,
    )
    normal = ParagraphStyle(
        "NormalL", fontName=base_font, fontSize=BODY_SIZE_PT,
        leading=leading, spaceAfter=PARA_SPACE_AFTER_PT, alignment=TA_LEFT
    )
    justify = ParagraphStyle("JustifyL", parent=normal, alignment=TA_JUSTIFY)
    cell_style = ParagraphStyle("Cell", parent=normal, spaceAfter=0, alignment=TA_LEFT)

    def two_col_table(left_html, right_html, second_row=None):
        """A borderless 2-column row, or 2 rows, using real fixed-width
        columns so alignment never depends on text length."""
        rows = [[Paragraph(left_html, cell_style), Paragraph(right_html, cell_style)]]
        if second_row:
            rows.append([Paragraph(second_row[0], cell_style), Paragraph(second_row[1], cell_style)])
        t = Table(rows, colWidths=[col_width * inch, (content_width - col_width) * inch])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        return t

    story = []

    if os.path.exists(LOGO_FILE):
        img = RLImage(LOGO_FILE, width=1.2 * inch, height=1.2 * inch)
        img.hAlign = "CENTER"
        story.append(img)
        story.append(Spacer(1, 6))

    story.append(Paragraph("Student Undertaking for Ethical Academic Practice", title_style))
    story.append(Paragraph("Vidyalankar Institute of Technology, Mumbai", title_style))
    story.append(Spacer(1, PARA_SPACE_AFTER_PT))

    meta_table_rows = [
        [Paragraph(f"Assignment No: <b>{data['assignment_no']}</b>", cell_style), Paragraph("", cell_style)],
        [Paragraph(f"Branch: <b>{data['branch']}</b>", cell_style), Paragraph(f"Vertical: <b>{data['vertical']}</b>", cell_style)],
        [Paragraph(f"Semester: <b>{data['semester']}</b>", cell_style), Paragraph(f"Division: <b>{data['division']}</b>", cell_style)],
        [Paragraph(f"Subject Name: <b>{data['subject_name']}</b>", cell_style), Paragraph(f"Subject Code: <b>{data['subject_code']}</b>", cell_style)],
    ]
    t_meta = Table(meta_table_rows, colWidths=[col_width * inch, (content_width - col_width) * inch])
    t_meta.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, PARA_SPACE_AFTER_PT))

    story.append(Paragraph(DECLARATION_INTRO.replace("\u2610", "&#9744;"), normal))

    option_lines = []
    for idx, text in enumerate(DECLARATION_OPTIONS, start=1):
        mark = "&#10003;" if idx in data["selected_options"] else "&#9744;"
        option_lines.append(f"{mark} ({idx}) {text}")
    story.append(Paragraph("<br/>".join(option_lines), normal))

    story.append(Paragraph(ETHICS_WARNING, justify))

    story.append(two_col_table(
        f"Student Name: <b>{data['student_name']}</b>",
        f"Roll No.: <b>{data['roll_no']}</b>",
    ))
    story.append(Spacer(1, 2))

    signature_path = data.get("signature_path")
    sig_label = "Signature:"
    date_label = f"Date: <b>{data['date']}</b>"
    if signature_path:
        processed_sig, tw, th = process_signature_image(signature_path)
        if processed_sig:
            if not getattr(processed_sig, "name", None):
                processed_sig.name = "signature.png"
            sig_img = RLImage(processed_sig, width=tw * inch, height=th * inch)
            sig_row = Table(
                [[Paragraph(sig_label, normal), sig_img, Paragraph(date_label, normal)]],
                colWidths=[0.85 * inch, tw * inch + 0.2 * inch,
                           (content_width - 0.85 - tw - 0.2) * inch],
            )
            sig_row.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]))
            story.append(sig_row)
        else:
            story.append(two_col_table(sig_label, date_label))
    else:
        story.append(two_col_table(sig_label, date_label))

    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        topMargin=1 * inch, bottomMargin=1 * inch,
        leftMargin=1 * inch, rightMargin=1 * inch,
    )
    doc.build(story)


def generate_docx_bytes(data):
    """Generates a DOCX file in memory using gendoc.py logic and returns a BytesIO buffer."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        generate_docx(data, tmp_path)
        with open(tmp_path, "rb") as f:
            buf = io.BytesIO(f.read())
        buf.seek(0)
        return buf
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


def generate_pdf_bytes(data):
    """Generates a PDF file in memory using gendoc.py logic and returns a BytesIO buffer."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        generate_pdf(data, tmp_path)
        with open(tmp_path, "rb") as f:
            buf = io.BytesIO(f.read())
        buf.seek(0)
        return buf
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


# --------------------------------------------------------------------------
# 4. FILE NAME HELPER
# --------------------------------------------------------------------------
def build_filename(subject_code, assignment_no, ext):
    safe_assignment = assignment_no.replace(" ", "").replace("/", "-")
    return f"{subject_code}-{safe_assignment}-Declaration Form.{ext}"


# --------------------------------------------------------------------------
# 5. TKINTER GUI
# --------------------------------------------------------------------------
class DeclarationFormApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Declaration Form Generator")
        self.geometry("600x750")
        self.minsize(520, 500)
        self.resizable(True, True)

        main_canvas = tk.Canvas(self, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=main_canvas.yview)
        container = ttk.Frame(main_canvas)

        container.bind(
            "<Configure>",
            lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all"))
        )
        window_id = main_canvas.create_window((0, 0), window=container, anchor="nw")

        def _on_canvas_configure(event):
            main_canvas.itemconfig(window_id, width=event.width)

        main_canvas.bind("<Configure>", _on_canvas_configure)
        main_canvas.configure(yscrollcommand=scrollbar.set)

        def _on_mousewheel(event):
            main_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        main_canvas.bind_all("<MouseWheel>", _on_mousewheel)

        scrollbar.pack(side="right", fill="y")
        main_canvas.pack(side="left", fill="both", expand=True)

        pad = {"padx": 14, "pady": 4}
        row = 0

        container.columnconfigure(1, weight=1)

        ttk.Label(container, text="Declaration Form Generator", font=("Segoe UI", 14, "bold")) \
            .grid(row=row, column=0, columnspan=2, pady=(12, 2)); row += 1

        ttk.Label(
            container,
            text="DOCX \u2192 forms/docs   \u2022   PDF \u2192 forms/pdfs",
            foreground="#555555",
        ).grid(row=row, column=0, columnspan=2, pady=(0, 8)); row += 1

        if not all(os.path.exists(p) for p in FONT_FILES.values()):
            ttk.Label(
                container,
                text="\u26a0 fonts/ folder incomplete -- outputs will use a fallback font",
                foreground="#b45309",
            ).grid(row=row, column=0, columnspan=2, pady=(0, 8)); row += 1

        self.all_subjects = []
        for vert, subs in VERTICALS.items():
            for name, code in subs:
                self.all_subjects.append((name, code, vert))

        ttk.Label(container, text="Subject:").grid(row=row, column=0, sticky="w", **pad)
        self.subject_var = tk.StringVar()
        self.subject_combo = ttk.Combobox(
            container, textvariable=self.subject_var,
            values=[s[0] for s in self.all_subjects], state="readonly")
        self.subject_combo.grid(row=row, column=1, sticky="ew", **pad)
        self.subject_combo.bind("<<ComboboxSelected>>", self.on_subject_change)
        row += 1

        ttk.Label(container, text="Subject Code:").grid(row=row, column=0, sticky="w", **pad)
        self.code_var = tk.StringVar()
        ttk.Entry(container, textvariable=self.code_var, state="readonly") \
            .grid(row=row, column=1, sticky="ew", **pad)
        row += 1

        ttk.Label(container, text="Vertical:").grid(row=row, column=0, sticky="w", **pad)
        self.vertical_var = tk.StringVar()
        ttk.Entry(container, textvariable=self.vertical_var, state="readonly") \
            .grid(row=row, column=1, sticky="ew", **pad)
        row += 1

        ttk.Label(container, text="Assignment / Experiment No:").grid(row=row, column=0, sticky="w", **pad)
        self.assignment_var = tk.StringVar(value="Experiment 1")
        ttk.Entry(container, textvariable=self.assignment_var) \
            .grid(row=row, column=1, sticky="ew", **pad)
        row += 1

        fields = [
            ("Student Name", "student_name", "Anirudh Ghanshyam Sarve"),
            ("Roll No.", "roll_no", "24102A0062"),
            ("Branch", "branch", "CMPN"),
            ("Semester", "semester", "V"),
            ("Division", "division", "A"),
        ]
        self.field_vars = {}
        for label, key, default in fields:
            ttk.Label(container, text=label + ":").grid(row=row, column=0, sticky="w", **pad)
            var = tk.StringVar(value=default)
            ttk.Entry(container, textvariable=var).grid(row=row, column=1, sticky="ew", **pad)
            self.field_vars[key] = var
            row += 1

        ttk.Label(container, text="Date:").grid(row=row, column=0, sticky="w", **pad)
        self.date_var = tk.StringVar(value=datetime.date.today().strftime("%d-%m-%Y"))
        ttk.Entry(container, textvariable=self.date_var).grid(row=row, column=1, sticky="ew", **pad)
        row += 1

        ttk.Label(container, text="Signature Image:").grid(row=row, column=0, sticky="w", **pad)
        sig_frame = ttk.Frame(container)
        sig_frame.grid(row=row, column=1, sticky="ew", **pad)
        sig_frame.columnconfigure(0, weight=1)
        self.signature_path_var = tk.StringVar(value="")
        ttk.Entry(sig_frame, textvariable=self.signature_path_var, state="readonly") \
            .grid(row=0, column=0, sticky="ew")
        ttk.Button(sig_frame, text="Browse...", command=self.browse_signature) \
            .grid(row=0, column=1, padx=(4, 0))
        ttk.Button(sig_frame, text="Clear", command=lambda: self.signature_path_var.set("")) \
            .grid(row=0, column=2, padx=(4, 0))
        row += 1

        ttk.Label(container, text="Tick applicable declaration statement(s):") \
            .grid(row=row, column=0, columnspan=2, sticky="w", padx=14, pady=(10, 2)); row += 1

        self.option_vars = []
        for idx, text in enumerate(DECLARATION_OPTIONS, start=1):
            var = tk.BooleanVar(value=(idx == 1))
            wrapped = self._wrap_text(f"({idx}) {text}", width=68)
            cb = ttk.Checkbutton(container, text=wrapped, variable=var)
            cb.grid(row=row, column=0, columnspan=2, sticky="w", padx=20, pady=2)
            self.option_vars.append(var)
            row += 1

        btn_frame = ttk.Frame(container)
        btn_frame.grid(row=row, column=0, columnspan=2, pady=16)
        ttk.Button(btn_frame, text="Generate DOCX", command=lambda: self.generate(["docx"])) \
            .grid(row=0, column=0, padx=6)
        ttk.Button(btn_frame, text="Generate PDF", command=lambda: self.generate(["pdf"])) \
            .grid(row=0, column=1, padx=6)
        ttk.Button(btn_frame, text="Generate Both", command=lambda: self.generate(["docx", "pdf"])) \
            .grid(row=0, column=2, padx=6)

        self.status_var = tk.StringVar(value="")
        ttk.Label(container, textvariable=self.status_var, foreground="green") \
            .grid(row=row + 1, column=0, columnspan=2, pady=(0, 14))

        if self.all_subjects:
            self.subject_combo.current(0)
            self.on_subject_change()

    @staticmethod
    def _wrap_text(text, width=68):
        import textwrap
        return "\n".join(textwrap.wrap(text, width=width))

    def browse_signature(self):
        path = filedialog.askopenfilename(
            title="Choose signature image",
            filetypes=[("Image files", "*.png *.jpg *.jpeg"), ("All files", "*.*")],
        )
        if path:
            self.signature_path_var.set(path)

    def on_subject_change(self, event=None):
        subject = self.subject_var.get()
        for name, code, vertical in self.all_subjects:
            if name == subject:
                self.code_var.set(code)
                self.vertical_var.set(vertical)
                break

    def collect_data(self):
        selected = [i + 1 for i, v in enumerate(self.option_vars) if v.get()]
        if not selected:
            messagebox.showwarning("Missing selection", "Please tick at least one declaration statement.")
            return None
        if not self.assignment_var.get().strip():
            messagebox.showwarning("Missing field", "Please enter the Assignment / Experiment No.")
            return None
        data = {
            "vertical": self.vertical_var.get(),
            "subject_name": self.subject_var.get(),
            "subject_code": self.code_var.get(),
            "assignment_no": self.assignment_var.get().strip(),
            "date": self.date_var.get().strip(),
            "selected_options": selected,
            "signature_path": self.signature_path_var.get().strip() or None,
        }
        for key, var in self.field_vars.items():
            data[key] = var.get().strip()
        return data

    def generate(self, formats):
        data = self.collect_data()
        if data is None:
            return

        generated = []
        try:
            if "docx" in formats:
                fname = build_filename(data["subject_code"], data["assignment_no"], "docx")
                path = os.path.join(DOCX_DIR, fname)
                generate_docx(data, path)
                generated.append(path)
            if "pdf" in formats:
                fname = build_filename(data["subject_code"], data["assignment_no"], "pdf")
                path = os.path.join(PDF_DIR, fname)
                generate_pdf(data, path)
                generated.append(path)
        except Exception as exc:
            messagebox.showerror("Error generating file", str(exc))
            return

        self.status_var.set("Saved: " + ", ".join(os.path.basename(p) for p in generated))
        messagebox.showinfo("Done", "File(s) generated:\n" + "\n".join(generated))


# --------------------------------------------------------------------------
# 6. CLI FALLBACK (in case Tkinter / a display is unavailable, e.g. servers)
# --------------------------------------------------------------------------
def run_cli():
    print("Declaration Form Generator (CLI mode)\n")
    all_subjects = []
    for vert, subs in VERTICALS.items():
        for name, code in subs:
            all_subjects.append((name, code, vert))

    for i, (name, code, vert) in enumerate(all_subjects, 1):
        print(f"  {i}. {name} ({code}) [{vert}]")
    s_idx = int(input("Choose subject number: ")) - 1
    subject_name, subject_code, vertical = all_subjects[s_idx]

    data = {
        "vertical": vertical,
        "subject_name": subject_name,
        "subject_code": subject_code,
        "assignment_no": input("Assignment / Experiment No: ") or "Experiment 1",
        "student_name": input("Student Name [Anirudh Ghanshyam Sarve]: ") or "Anirudh Ghanshyam Sarve",
        "roll_no": input("Roll No. [24102A0062]: ") or "24102A0062",
        "branch": input("Branch [CMPN]: ") or "CMPN",
        "semester": input("Semester [V]: ") or "V",
        "division": input("Division [A]: ") or "A",
        "date": input(f"Date [{datetime.date.today().strftime('%d-%m-%Y')}]: ")
        or datetime.date.today().strftime("%d-%m-%Y"),
        "signature_path": input("Path to signature image (leave blank to skip): ").strip() or None,
    }
    print("Tick applicable statement(s), comma separated (e.g. 1,3):")
    for idx, text in enumerate(DECLARATION_OPTIONS, start=1):
        print(f"  ({idx}) {text}")
    raw = input("> ")
    data["selected_options"] = [int(x) for x in raw.split(",") if x.strip().isdigit()] or [1]

    fmt = input("Generate [docx/pdf/both]: ").strip().lower() or "both"
    formats = ["docx", "pdf"] if fmt == "both" else [fmt]

    for f in formats:
        fname = build_filename(data["subject_code"], data["assignment_no"], f)
        if f == "docx":
            path = os.path.join(DOCX_DIR, fname)
            generate_docx(data, path)
        else:
            path = os.path.join(PDF_DIR, fname)
            generate_pdf(data, path)
        print("Saved:", path)


if __name__ == "__main__":
    if "--cli" in sys.argv:
        run_cli()
    else:
        try:
            app = DeclarationFormApp()
            app.mainloop()
        except tk.TclError:
            print("No display available for the GUI -- falling back to CLI mode.\n")
            run_cli()
