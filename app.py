import io
import os
import json
import datetime
import urllib.request
import streamlit as st
from PIL import Image

# --------------------------------------------------------------------------
# 1. VERTICAL / SUBJECT DATA
# --------------------------------------------------------------------------
VERTICALS = {
    "HSSM_VEC": [
        ("E-Waste and Environmental Management", "VEC02"),
    ],
    "MDC_MDM": [
        ("Intelligent Mobile Robotics", "MDMRB03"),
        ("Innovation Management and Scaling Startups", "MDMIE03"),
        ("Machine Learning Applications in Bioinfomatics", "MDMBI03"),
        ("Strategic Marketing and Business Planning", "MDMBI03"),
    ],
    "MDC_OEC": [
        ("Professional Competency Development", "OEC15"),
    ],
    "PC_PCC": [
        ("Artificial Intelligence", "PCCE10T"),
        ("Computer Networks", "PCCE11T"),
        ("Software Engineering", "PCCE12T"),
        ("Theory Of Computation", "PCCE09"),
    ],
    "PC_PEC": [
        ("Data Warehousing and Mining", "PECE02T"),
    ],
}

DECLARATION_OPTIONS = [
    "I have prepared the assignment / academic activity entirely by myself.",
    "I have completed the assignment / academic activity with academic guidance or discussion from seniors / friends / peers, while ensuring the work is my own.",
    "I have used AI-based tools only as an aid and have appropriately acknowledged or cited their use as per the ethical guidelines of the Institute.",
    "I have used AI-based tools and directly copy-pasted content for the assignment / academic activity.",
    "I have copied the assignment / academic activity from peers / friends.",
]

LOGO_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "college_logo.png")

ALL_SUBJECTS = []
for vert, subs in VERTICALS.items():
    for name, code in subs:
        ALL_SUBJECTS.append({"name": name, "code": code, "vertical": vert})

# --------------------------------------------------------------------------
# 2. SUPABASE REST HELPERS (Zero Extra Dependencies)
# --------------------------------------------------------------------------
def fetch_supabase_profile(url, key, roll_no):
    if not url or not key or not roll_no:
        return None
    try:
        endpoint = f"{url.rstrip('/')}/rest/v1/student_profiles?roll_no=eq.{roll_no}&select=*"
        req = urllib.request.Request(endpoint, headers={
            "apikey": key,
            "Authorization": f"Bearer {key}"
        })
        with urllib.request.urlopen(req, timeout=4) as resp:
            data = json.loads(resp.read().decode())
            return data[0] if data else None
    except Exception:
        return None

def upsert_supabase_profile(url, key, profile):
    if not url or not key or not profile.get("roll_no"):
        return
    try:
        endpoint = f"{url.rstrip('/')}/rest/v1/student_profiles"
        payload = json.dumps([profile]).encode()
        req = urllib.request.Request(endpoint, data=payload, headers={
            "apikey": key,
            "Authorization": f"Bearer {key}",
            "Content-Type": "application/json",
            "Prefer": "resolution=merge-duplicates"
        }, method="POST")
        urllib.request.urlopen(req, timeout=4)
    except Exception:
        pass

# --------------------------------------------------------------------------
# 3. SIGNATURE PROCESSOR (Fixed Ratio & Max Dimensions)
# --------------------------------------------------------------------------
def process_signature_image(img_bytes):
    if not img_bytes:
        return None, 0, 0
    try:
        from PIL import ImageEnhance, ImageFilter
        img_bytes.seek(0)
        img = Image.open(img_bytes).convert("RGBA")

        w_orig, h_orig = img.size
        if w_orig > 0 and h_orig > 0:
            scale = 3
            img_work = img.resize((w_orig * scale, h_orig * scale), Image.Resampling.LANCZOS)
        else:
            img_work = img

        gray = img_work.convert("L")
        gray_enhanced = ImageEnhance.Contrast(gray).enhance(2.2)

        mask = Image.eval(gray_enhanced, lambda v: 255 - v if v < 225 else 0)
        mask_smooth = mask.filter(ImageFilter.GaussianBlur(radius=1.2))
        mask_final = Image.eval(mask_smooth, lambda v: min(255, int(v * 1.6)))

        ink_layer = Image.new("RGBA", img_work.size, (15, 23, 42, 255))
        ink_layer.putalpha(mask_final)

        smooth_img = ink_layer.resize((w_orig, h_orig), Image.Resampling.LANCZOS)

        bbox = smooth_img.getchannel('A').getbbox()
        if bbox:
            smooth_img = smooth_img.crop(bbox)

        out = io.BytesIO()
        smooth_img.save(out, format="PNG")
        out.name = "signature.png"
        out.seek(0)

        w, h = smooth_img.size
        aspect = w / float(h) if h > 0 else 2.4

        MAX_H = 0.32
        MAX_W = 0.80

        target_h = MAX_H
        target_w = target_h * aspect

        if target_w > MAX_W:
            target_w = MAX_W
            target_h = target_w / aspect

        return out, target_w, target_h
    except Exception:
        img_bytes.seek(0)
        if not hasattr(img_bytes, 'name'):
            img_bytes.name = "signature.png"
        return img_bytes, 0.79, 0.32

# --------------------------------------------------------------------------
# 4. IN-MEMORY DOCX GENERATION
# --------------------------------------------------------------------------
def generate_docx_bytes(data):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    if os.path.exists(LOGO_FILE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(LOGO_FILE, width=Inches(1.2))

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)

    def add_line(pairs):
        p = doc.add_paragraph()
        for i, (label, value) in enumerate(pairs):
            if i > 0:
                p.add_run("\u2003\u2003\u2003\u2003")
            p.add_run(label)
            r = p.add_run(str(value))
            r.bold = True
        return p

    add_title("Student Undertaking for Ethical Academic Practice")
    add_title("Vidyalankar Institute of Technology, Mumbai")
    doc.add_paragraph("")

    add_line([("Assignment No: ", data["assignment_no"])])
    add_line([("Branch: ", data["branch"]), ("Vertical: ", data["vertical"])])

    p = doc.add_paragraph()
    p.add_run("Semester: ")
    p.add_run(data["semester"]).bold = True
    p.add_run("\u2003\u2003\u2003\u2003Division: ")
    p.add_run(data["division"]).bold = True
    p.add_run("\nSubject Name: ")
    p.add_run(data["subject_name"]).bold = True
    p.add_run("\u2003\u2003\u2003\u2003Subject Code: ")
    p.add_run(data["subject_code"]).bold = True

    doc.add_paragraph(
        "I hereby declare that for the assignment / academic activity "
        "submitted by me, I confirm the following statement(s) by ticking "
        "(\u2611) the appropriate box(es):"
    )

    for idx, text in enumerate(DECLARATION_OPTIONS, start=1):
        mark = "\u2611" if idx in data["selected_options"] else "\u2610"
        doc.add_paragraph(f"{mark} ({idx}) {text}")

    warn = doc.add_paragraph(
        "I understand that selecting options (4) or (5) indicates unethical "
        "academic practice and may attract academic penalties, including "
        "rejection of the assignment or disciplinary action, as per the "
        "rules of Vidyalankar Institute of Technology, Mumbai. I submit "
        "this declaration truthfully and accept full responsibility for "
        "the same."
    )
    warn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph()
    p.add_run("Student Name: ")
    p.add_run(data["student_name"]).bold = True
    p.add_run("\u2003\u2003Roll No.: ")
    p.add_run(data["roll_no"]).bold = True

    sig_p = doc.add_paragraph()
    sig_p.add_run("Signature: ")
    
    sig_stream = data.get("signature_bytes")
    if sig_stream:
        processed_sig, tw, th = process_signature_image(sig_stream)
        if processed_sig:
            sig_p.add_run().add_picture(processed_sig, width=Inches(tw), height=Inches(th))
        else:
            sig_p.add_run("\u2003" * 10)
    else:
        sig_p.add_run("\u2003" * 10)
    sig_p.add_run("\u2003\u2003Date: ")
    sig_p.add_run(data["date"]).bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --------------------------------------------------------------------------
# 5. IN-MEMORY PDF GENERATION
# --------------------------------------------------------------------------
def generate_pdf_bytes(data):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
    from reportlab.lib.utils import ImageReader

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleC", parent=styles["Normal"], alignment=TA_CENTER,
        fontName="Helvetica-Bold", fontSize=14, spaceAfter=4,
    )
    normal = ParagraphStyle("NormalL", parent=styles["Normal"], fontSize=11, spaceAfter=8, leading=15)
    justify = ParagraphStyle("JustifyL", parent=normal, alignment=TA_JUSTIFY)

    story = []

    if os.path.exists(LOGO_FILE):
        img = RLImage(LOGO_FILE, width=1.2 * inch, height=1.2 * inch)
        img.hAlign = "CENTER"
        story.append(img)
        story.append(Spacer(1, 6))

    story.append(Paragraph("Student Undertaking for Ethical Academic Practice", title_style))
    story.append(Paragraph("Vidyalankar Institute of Technology, Mumbai", title_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph(f"Assignment No: <b>{data['assignment_no']}</b>", normal))
    story.append(Paragraph(
        f"Branch: <b>{data['branch']}</b>"
        f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
        f"Vertical: <b>{data['vertical']}</b>", normal))
    story.append(Paragraph(
        f"Semester: <b>{data['semester']}</b>"
        f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
        f"Division: <b>{data['division']}</b>", normal))
    story.append(Paragraph(
        f"Subject Name: <b>{data['subject_name']}</b>"
        f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
        f"Subject Code: <b>{data['subject_code']}</b>", normal))

    story.append(Paragraph(
        "I hereby declare that for the assignment / academic activity "
        "submitted by me, I confirm the following statement(s) by ticking "
        "(&#9745;) the appropriate box(es):", normal))

    for idx, text in enumerate(DECLARATION_OPTIONS, start=1):
        mark = "&#9745;" if idx in data["selected_options"] else "&#9744;"
        story.append(Paragraph(f"{mark} ({idx}) {text}", normal))

    story.append(Paragraph(
        "I understand that selecting options (4) or (5) indicates unethical "
        "academic practice and may attract academic penalties, including "
        "rejection of the assignment or disciplinary action, as per the "
        "rules of Vidyalankar Institute of Technology, Mumbai. I submit "
        "this declaration truthfully and accept full responsibility for "
        "the same.", justify))

    story.append(Spacer(1, 14))
    story.append(Paragraph(
        f"Student Name: <b>{data['student_name']}</b>"
        f"&nbsp;&nbsp;&nbsp;&nbsp;Roll No.: <b>{data['roll_no']}</b>", normal))

    sig_bytes = data.get("signature_bytes")
    if sig_bytes:
        processed_sig, tw, th = process_signature_image(sig_bytes)
        if processed_sig:
            if not getattr(processed_sig, 'name', None):
                processed_sig.name = "signature.png"
            sig_img = RLImage(processed_sig, width=tw * inch, height=th * inch)
            sig_row = Table(
                [[Paragraph("Signature:", normal), sig_img, Paragraph(f"Date: <b>{data['date']}</b>", normal)]],
                colWidths=[0.9 * inch, tw * inch + 0.2 * inch, 2.5 * inch],
            )
            sig_row.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]))
            story.append(sig_row)
        else:
            story.append(Paragraph(
                f"Signature&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                f"Date: <b>{data['date']}</b>", normal))
    else:
        story.append(Paragraph(
            f"Signature&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
            f"Date: <b>{data['date']}</b>", normal))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        leftMargin=1 * inch, rightMargin=1 * inch,
    )
    doc.build(story)
    buffer.seek(0)
    return buffer

# --------------------------------------------------------------------------
# 6. STREAMLIT WEB UI
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="VIT Declaration Form Generator",
    page_icon="📜",
    layout="centered"
)

# Custom Styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Plus Jakarta Sans', sans-serif;
    }
    
    .stApp {
        background: linear-gradient(135deg, #0F172A 0%, #1E1B4B 50%, #0F172A 100%);
    }

    .main .block-container {
        max-width: 780px;
        padding-top: 2rem;
        padding-bottom: 3rem;
    }

    /* Force all labels, text and headings to bright white/slate */
    label, .stWidgetLabel, [data-testid="stWidgetLabel"] p, [data-testid="stMarkdownContainer"] p {
        color: #F8FAFC !important;
        font-weight: 600 !important;
    }

    h1, h2, h3, h4, h5, h6, .stSubheader {
        color: #F8FAFC !important;
        font-weight: 700 !important;
    }

    .vit-badge {
        display: inline-block;
        background: rgba(99, 102, 241, 0.2);
        color: #A5B4FC;
        font-weight: 700;
        font-size: 0.8rem;
        padding: 0.35rem 0.9rem;
        border-radius: 50px;
        letter-spacing: 0.5px;
        margin-bottom: 0.5rem;
    }

    .main-title {
        color: #FFFFFF;
        font-weight: 800;
        font-size: 2rem;
        letter-spacing: -0.5px;
        margin-bottom: 0.2rem;
    }

    .sub-title {
        color: #CBD5E1;
        font-size: 0.95rem;
        margin-bottom: 1.5rem;
    }

    /* Inputs background & text styling */
    input[type="text"], div[data-baseweb="select"] > div {
        background-color: #F8FAFC !important;
        color: #0F172A !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
    }
</style>
""", unsafe_allow_html=True)

header_col1, header_col2 = st.columns([0.82, 0.18])
with header_col1:
    st.markdown('<div class="vit-badge">🎓 VIDYALANKAR INSTITUTE OF TECHNOLOGY</div>', unsafe_allow_html=True)
    st.markdown('<div class="main-title">Declaration Form Generator</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-title">Generate and download official declaration forms (.docx & .pdf) formatted for VIT submissions.</div>', unsafe_allow_html=True)

with header_col2:
    with st.popover("⚙️", help="Customise Details & Add Subjects"):
        st.markdown("### ⚙️ Customise Profile & Subjects")
        
        # Section 1: Add Custom Subject
        st.caption("➕ **Add Custom Subject**")
        c_name = st.text_input("Subject Name", key="pop_c_name", placeholder="e.g. Robotics")
        c_code = st.text_input("Subject Code", key="pop_c_code", placeholder="e.g. ROB101")
        c_vert = st.text_input("Vertical", value="CUSTOM", key="pop_c_vert")
        
        if st.button("➕ Add Subject", key="btn_add_subj"):
            if c_name and c_code:
                if "custom_subjects" not in st.session_state:
                    st.session_state["custom_subjects"] = []
                st.session_state["custom_subjects"].append({"name": c_name, "code": c_code, "vertical": c_vert})
                st.toast(f"✅ Added subject: {c_name} ({c_code})")
                st.rerun()
            else:
                st.warning("Enter both Subject Name and Code.")

        st.markdown("---")
        
        # Section 2: Customise Profile Defaults
        st.caption("👤 **Customise Profile Defaults**")
        p_name = st.text_input("Full Name", value=st.session_state.get("student_name", "Anirudh Ghanshyam Sarve"), key="pop_p_name")
        p_branch = st.text_input("Branch", value=st.session_state.get("branch", "CMPN"), key="pop_p_branch")
        p_sem = st.text_input("Semester", value=st.session_state.get("semester", "V"), key="pop_p_sem")
        p_div = st.text_input("Division", value=st.session_state.get("division", "A"), key="pop_p_div")
        
        if st.button("💾 Save Profile Defaults", key="btn_save_prof"):
            st.session_state["student_name"] = p_name
            st.session_state["branch"] = p_branch
            st.session_state["semester"] = p_sem
            st.session_state["division"] = p_div
            
            # Sync to Supabase if configured
            sp_url = st.session_state.get("sb_url")
            sp_key = st.session_state.get("sb_key")
            curr_roll = st.session_state.get("roll_no", "24102A0062")
            if sp_url and sp_key and curr_roll:
                upsert_supabase_profile(sp_url, sp_key, {
                    "roll_no": curr_roll,
                    "student_name": p_name,
                    "branch": p_branch,
                    "semester": p_sem,
                    "division": p_div,
                    "custom_subjects": json.dumps(st.session_state.get("custom_subjects", []))
                })
            st.toast("✅ Profile defaults saved and synced!")
            st.rerun()

# Supabase Config load from secrets / env
env_sb_url = os.environ.get("SUPABASE_URL", "")
env_sb_key = os.environ.get("SUPABASE_ANON_KEY", "")
try:
    if "SUPABASE_URL" in st.secrets:
        env_sb_url = st.secrets["SUPABASE_URL"]
    if "SUPABASE_ANON_KEY" in st.secrets:
        env_sb_key = st.secrets["SUPABASE_ANON_KEY"]
except Exception:
    pass

if env_sb_url and env_sb_key:
    st.session_state["sb_url"] = env_sb_url
    st.session_state["sb_key"] = env_sb_key

# Subject selection
with st.container():
    st.subheader("📚 Subject & Academic Details")
    all_subj_list = ALL_SUBJECTS + st.session_state.get("custom_subjects", [])
    subject_names = [s["name"] for s in all_subj_list]
    selected_subject_name = st.selectbox("Select Subject", options=subject_names)
    selected_subj_info = next((s for s in all_subj_list if s["name"] == selected_subject_name), ALL_SUBJECTS[0])

    col_sub1, col_sub2 = st.columns(2)
    with col_sub1:
        st.text_input("Subject Code", value=selected_subj_info["code"], disabled=True)
    with col_sub2:
        st.text_input("Vertical", value=selected_subj_info["vertical"], disabled=True)

    assignment_no = st.text_input("Assignment / Experiment No", value="Experiment 1")

st.markdown("---")

# Student Details
with st.container():
    st.subheader("👤 Student Details")
    
    # Roll number auto-fetch trigger
    roll_input = st.text_input("Roll No. (Enter to fetch profile)", value=st.session_state.get("roll_no", "24102A0062"), key="roll_input_field")
    
    if roll_input != st.session_state.get("last_roll"):
        st.session_state["last_roll"] = roll_input
        # Fetch from Supabase if configured
        sp_url = st.session_state.get("sb_url")
        sp_key = st.session_state.get("sb_key")
        if sp_url and sp_key:
            remote = fetch_supabase_profile(sp_url, sp_key, roll_input.strip())
            if remote:
                st.session_state["student_name"] = remote.get("student_name", "")
                st.session_state["branch"] = remote.get("branch", "")
                st.session_state["semester"] = remote.get("semester", "")
                st.session_state["division"] = remote.get("division", "")
                if remote.get("custom_subjects"):
                    try:
                        c_sub = json.loads(remote["custom_subjects"]) if isinstance(remote["custom_subjects"], str) else remote["custom_subjects"]
                        if isinstance(c_sub, list):
                            st.session_state["custom_subjects"] = c_sub
                    except Exception:
                        pass
                st.toast(f"⚡ Synced profile for {roll_input} from Supabase!")

    col1, col2 = st.columns(2)
    with col1:
        student_name = st.text_input("Student Name", value=st.session_state.get("student_name", "Anirudh Ghanshyam Sarve"))
        branch = st.text_input("Branch", value=st.session_state.get("branch", "CMPN"))
        division = st.text_input("Division", value=st.session_state.get("division", "A"))
    with col2:
        roll_no = roll_input
        semester = st.text_input("Semester", value=st.session_state.get("semester", "V"))
        date_str = st.text_input("Date (DD-MM-YYYY)", value=datetime.date.today().strftime("%d-%m-%Y"))

st.markdown("---")

# Signature Section
with st.container():
    st.subheader("✍️ Signature Input")
    sig_method = st.radio("Choose Signature Method:", ["Draw Signature", "Upload Image File", "Skip"], horizontal=True)

    sig_bytes_io = None

    if sig_method == "Draw Signature":
        try:
            from streamlit_drawable_canvas import st_canvas
            st.caption("Use your mouse or touch pad to sign inside the box below:")
            canvas_result = st_canvas(
                fill_color="rgba(255, 255, 255, 0)",
                stroke_width=3,
                stroke_color="#0F172A",
                background_color="#FFFFFF",
                height=140,
                width=500,
                drawing_mode="freedraw",
                key="canvas_sig",
            )
            if canvas_result.image_data is not None:
                img = Image.fromarray(canvas_result.image_data.astype('uint8'), 'RGBA')
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)
                sig_bytes_io = buf
        except Exception as e:
            st.error(f"Canvas error: {e}")

    elif sig_method == "Upload Image File":
        uploaded_sig = st.file_uploader("Upload Signature (PNG / JPG)", type=["png", "jpg", "jpeg"])
        if uploaded_sig is not None:
            sig_bytes_io = io.BytesIO(uploaded_sig.getvalue())
            st.image(uploaded_sig, caption="Uploaded Signature Preview", width=180)

st.markdown("---")

# Declaration checkboxes
with st.container():
    st.subheader("✅ Ethical Declaration Statements")
    st.write("Tick applicable statement(s):")

    selected_options = []
    for idx, option_text in enumerate(DECLARATION_OPTIONS, start=1):
        checked = st.checkbox(f"({idx}) {option_text}", value=(idx == 1), key=f"opt_{idx}")
        if checked:
            selected_options.append(idx)

st.markdown("---")

# Generate buttons
if not selected_options:
    st.warning("⚠️ Please select at least one declaration statement above.")
else:
    form_data = {
        "vertical": selected_subj_info["vertical"],
        "subject_name": selected_subj_info["name"],
        "subject_code": selected_subj_info["code"],
        "assignment_no": assignment_no.strip(),
        "student_name": student_name.strip(),
        "roll_no": roll_no.strip(),
        "branch": branch.strip(),
        "semester": semester.strip(),
        "division": division.strip(),
        "date": date_str.strip(),
        "selected_options": selected_options,
        "signature_bytes": sig_bytes_io,
    }

    # Save to Supabase background task
    sp_url = st.session_state.get("sb_url")
    sp_key = st.session_state.get("sb_key")
    if sp_url and sp_key:
        upsert_supabase_profile(sp_url, sp_key, {
            "roll_no": roll_no.strip(),
            "student_name": student_name.strip(),
            "branch": branch.strip(),
            "semester": semester.strip(),
            "division": division.strip(),
            "last_subject_code": selected_subj_info["code"],
            "last_subject_name": selected_subj_info["name"],
            "updated_at": datetime.datetime.now().isoformat()
        })

    safe_assignment = assignment_no.strip().replace(" ", "").replace("/", "-")
    base_filename = f"{selected_subj_info['code']}-{safe_assignment}-Declaration Form"

    st.subheader("📥 Download Generated Declaration Form")
    col_dl1, col_dl2 = st.columns(2)

    with col_dl1:
        docx_buffer = generate_docx_bytes(form_data)
        st.download_button(
            label="📄 Download Word (.docx)",
            data=docx_buffer,
            file_name=f"{base_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with col_dl2:
        pdf_buffer = generate_pdf_bytes(form_data)
        st.download_button(
            label="📕 Download PDF (.pdf)",
            data=pdf_buffer,
            file_name=f"{base_filename}.pdf",
            mime="application/pdf",
            use_container_width=True
        )
